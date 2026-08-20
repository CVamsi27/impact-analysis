#!/usr/bin/env python3
"""
Create a non-technical Word report from the analysis outputs.

This script intentionally contains ONLY end-user friendly results and avoids
implementation details.
"""

from __future__ import annotations

from pathlib import Path


def _pct_increase(group_value: float, baseline_value: float) -> float:
    if baseline_value == 0:
        return 0.0
    return (group_value - baseline_value) / baseline_value * 100.0


def main() -> None:
    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise SystemExit(
            "Missing dependency: python-docx\n"
            "Run:\n"
            "  source .venv/bin/activate\n"
            "  python -m pip install python-docx\n"
        ) from exc

    output_dir = Path("analysis_outputs")
    output_dir.mkdir(parents=True, exist_ok=True)
    out_path = output_dir / "Second_Victim_Survey_Results_(Non-Technical).docx"

    # Key overall means from analysis_outputs/scale_summary.csv (already computed).
    overall = {
        "PHQ9": 4.3876,
        "GAD7": 4.2121,
        "PerformanceImpact": 2.3636,
        "WorkWithdrawal": 2.2189,
        "IntentLeave": 1.8207,
    }

    # Highlighted segments (from analysis_outputs/high_risk_segments.csv).
    segments = {
        "Designation": {
            "Post Graduate": {"GAD7": 6.4545},
            "Nursing Officer": {"PHQ9": 5.3333},
            "Nurse": {"PHQ9": 5.1818, "PerformanceImpact": 3.0455, "WorkWithdrawal": 2.6061},
        },
        "Specialization": {
            "Ophthalmology": {"PHQ9": 7.4, "GAD7": 5.6},
            "Pediatrics": {"PHQ9": 5.1719, "GAD7": 5.125},
            "Obstetrics and Gynaecology": {"PHQ9": 5.2727},
        },
    }

    # Relationships (from analysis_outputs/relationship_insights.csv).
    relationship_points = [
        "When intention to leave is higher, work-withdrawal (mental health day / distraction) also tends to be higher.",
        "Depression symptoms (PHQ‑9) and anxiety symptoms (GAD‑7) tend to rise together.",
        "When second‑victim emotional impact is higher, work performance impact tends to be higher.",
        "When work performance impact is higher, intention to leave tends to be higher.",
        "When second‑victim emotional impact is higher, the need for support resources tends to be higher.",
    ]

    regression_takeaways = [
        "Second‑victim emotional impact is linked with higher performance impact (even after accounting for other measured factors).",
        "Second‑victim emotional impact is linked with higher anxiety and higher burnout (after adjustment).",
        "Higher performance impact is linked with higher intention to leave (after adjustment).",
        "More organisational support is linked with lower intention to leave (after adjustment).",
        "More supervisor support is linked with lower performance impact (after adjustment).",
    ]

    doc = Document()

    doc.add_heading("Second Victim Phenomenon Survey — Results (Non‑Technical)", level=1)
    p = doc.add_paragraph(
        "This report summarizes the main patterns found in the survey responses in simple terms. "
        "It describes relationships and group differences observed in the data."
    )
    p.paragraph_format.space_after = 8

    doc.add_heading("1) Key relationships (what tends to increase together)", level=2)
    for item in relationship_points:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2) Adjusted findings (regression-style takeaways)", level=2)
    doc.add_paragraph(
        "These points describe patterns that still appear after considering other variables together. "
        "They are still associations (not proof of cause)."
    )
    for item in regression_takeaways:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3) Which designation seems more affected (higher averages)", level=2)
    doc.add_paragraph(
        "Below are the designation groups that stand out for higher average scores. "
        "Percentages compare the group’s average to the overall average."
    )

    for designation, values in segments["Designation"].items():
        lines = []
        if "GAD7" in values:
            pct = _pct_increase(values["GAD7"], overall["GAD7"])
            lines.append(f"Anxiety (GAD‑7): {values['GAD7']:.2f} vs {overall['GAD7']:.2f} overall (about {pct:.0f}% higher)")
        if "PHQ9" in values:
            pct = _pct_increase(values["PHQ9"], overall["PHQ9"])
            lines.append(f"Depression (PHQ‑9): {values['PHQ9']:.2f} vs {overall['PHQ9']:.2f} overall (about {pct:.0f}% higher)")
        if "PerformanceImpact" in values:
            pct = _pct_increase(values["PerformanceImpact"], overall["PerformanceImpact"])
            lines.append(
                f"Performance impact: {values['PerformanceImpact']:.2f} vs {overall['PerformanceImpact']:.2f} overall (about {pct:.0f}% higher)"
            )
        if "WorkWithdrawal" in values:
            pct = _pct_increase(values["WorkWithdrawal"], overall["WorkWithdrawal"])
            lines.append(
                f"Work‑withdrawal impact: {values['WorkWithdrawal']:.2f} vs {overall['WorkWithdrawal']:.2f} overall (about {pct:.0f}% higher)"
            )

        doc.add_paragraph(f"{designation}", style="List Bullet")
        for line in lines:
            doc.add_paragraph(line, style="List Bullet 2")

    doc.add_heading("4) Which specialization seems more affected (higher averages)", level=2)
    doc.add_paragraph(
        "Below are the specialization groups that stand out for higher average scores. "
        "If a group has a small number of respondents, treat it as an early signal rather than a final conclusion."
    )

    for spec, values in segments["Specialization"].items():
        doc.add_paragraph(f"{spec}", style="List Bullet")
        if "PHQ9" in values:
            pct = _pct_increase(values["PHQ9"], overall["PHQ9"])
            doc.add_paragraph(
                f"Depression (PHQ‑9): {values['PHQ9']:.2f} vs {overall['PHQ9']:.2f} overall (about {pct:.0f}% higher)",
                style="List Bullet 2",
            )
        if "GAD7" in values:
            pct = _pct_increase(values["GAD7"], overall["GAD7"])
            doc.add_paragraph(
                f"Anxiety (GAD‑7): {values['GAD7']:.2f} vs {overall['GAD7']:.2f} overall (about {pct:.0f}% higher)",
                style="List Bullet 2",
            )

    doc.add_heading("5) Bottom-line summary (one paragraph)", level=2)
    doc.add_paragraph(
        "Overall, higher second‑victim emotional impact is linked with higher anxiety, depression, burnout, and work‑performance impact. "
        "Among designations, Post Graduates stand out most clearly for higher anxiety, while Nurses and Nursing Officers show higher depression and work‑impact signals. "
        "Among specializations, Ophthalmology shows the highest depression average in this dataset, and Pediatrics and Obstetrics & Gynaecology also show elevated depression/anxiety compared with the overall average."
    )

    doc.save(out_path)
    print(f"Created: {out_path}")


if __name__ == "__main__":
    main()

